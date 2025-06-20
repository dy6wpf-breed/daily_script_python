import os
import fitz
from docx import Document
from docx.shared import Inches
from tqdm import tqdm
from PIL import Image
import io
from deep_translator import GoogleTranslator
import time

def extract_pdf_content(pdf_path):
    """提取PDF的文本内容和图片"""
    text_content = []
    images = []
    
    doc = fitz.open(pdf_path)
    
    for page_num in range(len(doc)):
        page = doc[page_num]
        text_content.append(page.get_text())
        
        # 提取图片
        for img_index, img in enumerate(page.get_images(full=True)):
            try:
                xref = img[0]
                base_image = doc.extract_image(xref)
                image_bytes = base_image["image"]
                
                # 使用PIL验证和转换图片
                pil_image = Image.open(io.BytesIO(image_bytes))
                
                # 保存图片
                image_path = os.path.join(os.path.dirname(pdf_path), f'page_{page_num+1}_img_{img_index}.png')
                pil_image.save(image_path)
                
                images.append({
                    'path': image_path, 
                    'page_num': page_num
                })
            except Exception as e:
                print(f"图片提取失败：{e}")
    
    return text_content, images

def translate_text(text):
    """使用Google翻译API翻译文本"""
    try:
        # 如果文本为空，直接返回
        if not text.strip():
            return ""

        # 文本预处理
        text = text.replace('\n', ' ').replace('\r', ' ')
        text = ' '.join(text.split())
        
        # 如果处理后文本为空，直接返回
        if not text.strip():
            return ""

        # 将文本分成更小的块（每块最多1000字符）
        chunks = [text[i:i+1000] for i in range(0, len(text), 1000)]
        translated_chunks = []

        for chunk in chunks:
            if not chunk.strip():
                continue

            # 增加重试机制
            max_retries = 5
            for attempt in range(max_retries):
                try:
                    # 使用更保守的设置
                    translator = GoogleTranslator(source='auto', target='zh-CN')
                    translation = translator.translate(text=chunk)
                    
                    if translation:
                        translated_chunks.append(translation)
                        time.sleep(2)  # 每次翻译后等待2秒
                        break
                    else:
                        print(f"翻译返回空结果，正在重试...")
                        time.sleep(3)
                        
                except Exception as e:
                    print(f"翻译出错 (尝试 {attempt + 1}/{max_retries}): {str(e)}")
                    if attempt < max_retries - 1:
                        time.sleep(5)  # 出错后等待更长时间
                    else:
                        translated_chunks.append(chunk)  # 如果所有重试都失败，使用原文

        # 合并翻译结果
        result = ' '.join(translated_chunks)
        return result if result else text

    except Exception as e:
        print(f"翻译过程发生错误: {str(e)}")
        return text

def clean_text(text):
    """清理文本，移除不兼容的字符"""
    # 移除控制字符，但保留换行和制表符
    return ''.join(char for char in text if char >= ' ' or char in ['\n', '\t'])

def create_translated_docx(text_content, images, output_path):
    """创建翻译后的Word文档"""
    doc = Document()
    
    # 设置文档样式
    doc.styles['Normal'].font.name = 'SimSun'
    doc.styles['Normal'].font.size = 140000
    
    for page_num, page_text in tqdm(enumerate(text_content), desc="翻译和生成文档"):
        print(f"\n正在处理第 {page_num + 1} 页...")
        
        try:
            # 翻译文本
            translated_text = translate_text(page_text)
            
            # 清理文本
            translated_text = clean_text(translated_text)
            
            # 添加标题
            heading = doc.add_heading(f'第 {page_num + 1} 页', level=1)
            heading.style.font.name = 'SimHei'
            
            # 添加原文
            doc.add_paragraph("原文：").bold = True
            doc.add_paragraph(page_text)
            
            # 添加翻译
            doc.add_paragraph("翻译：").bold = True
            
            # 添加文本，保持段落格式
            paragraphs = translated_text.split('\n')
            for para in paragraphs:
                if para.strip():
                    try:
                        p = doc.add_paragraph(para.strip())
                        p.paragraph_format.line_spacing = 1.5
                        p.paragraph_format.space_after = 140000
                    except Exception as e:
                        print(f"添加段落失败: {e}")
                        continue
            
            # 添加该页的图片
            page_images = [img for img in images if img['page_num'] == page_num]
            for img in page_images:
                try:
                    doc.add_picture(img['path'], width=Inches(6.0))
                except Exception as e:
                    print(f"添加图片失败：{img['path']}, 错误：{e}")
            
            doc.add_page_break()
            
            # 定期保存文档
            if (page_num + 1) % 5 == 0:
                try:
                    doc.save(output_path)
                    print(f"已保存前 {page_num + 1} 页的翻译结果")
                except Exception as e:
                    print(f"保存文档时发生错误: {e}")
                    
        except Exception as e:
            print(f"处理第 {page_num + 1} 页时发生错误: {e}")
            continue
    
    # 最终保存文档
    try:
        doc.save(output_path)
    except Exception as e:
        print(f"保存最终文档时发生错误: {e}")

def main():
    # PDF输入路径
    pdf_path = r"E:\desk\25-15-2953.pdf"
    
    # 添加时间戳到输出文件名
    timestamp = time.strftime("%Y%m%d_%H%M%S")
    output_docx_path = rf'E:\desk\translated_literature_{timestamp}.docx'
    
    # 提取PDF内容
    text_content, images = extract_pdf_content(pdf_path)
    
    # 创建翻译文档
    create_translated_docx(text_content, images, output_docx_path)
    
    print(f"翻译完成，文档已保存至 {output_docx_path}")

if __name__ == "__main__":
    main()